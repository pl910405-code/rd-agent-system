"""
文档解析模块 - 支持PDF/Word/Excel/TXT文件解析与分块
"""
import os
import re
from typing import List, Dict, Optional
import fitz  # PyMuPDF
from docx import Document
from openpyxl import load_workbook


class DocumentProcessor:
    """文档解析与分块处理器"""

    SUPPORTED_TYPES = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".txt", ".csv", ".md"}

    def process_file(self, file_path: str) -> Dict:
        """解析文件，返回结构化文本内容"""
        ext = os.path.splitext(file_path)[1].lower()
        filename = os.path.basename(file_path)

        if ext == ".pdf":
            text, pages = self._process_pdf(file_path)
        elif ext in (".docx", ".doc"):
            text, pages = self._process_docx(file_path)
        elif ext in (".xlsx", ".xls"):
            text, pages = self._process_xlsx(file_path)
        elif ext in (".txt", ".csv", ".md"):
            text, pages = self._process_text(file_path)
        else:
            return {"error": f"不支持的文件类型: {ext}"}

        chunks = self.chunk_text(text, filename)
        return {
            "filename": filename,
            "file_path": file_path,
            "file_type": ext,
            "total_text": text,
            "total_chars": len(text),
            "page_count": len(pages) if pages else 1,
            "chunks": chunks,
            "pages": pages
        }

    def _process_pdf(self, file_path: str) -> tuple:
        """解析PDF文件"""
        text = ""
        pages = []
        try:
            doc = fitz.open(file_path)
            for i, page in enumerate(doc):
                page_text = page.get_text()
                pages.append({"page": i + 1, "text": page_text})
                text += page_text + "\n"
            doc.close()
        except Exception as e:
            text = f"PDF解析错误: {str(e)}"
        return text, pages

    def _process_docx(self, file_path: str) -> tuple:
        """解析Word文档"""
        text = ""
        pages = []
        try:
            doc = Document(file_path)
            current_page = []
            for para in doc.paragraphs:
                if para.text.strip():
                    current_page.append(para.text)
                    # 简单分页：每50段一页
                    if len(current_page) >= 50:
                        pages.append({"page": len(pages) + 1, "text": "\n".join(current_page)})
                        current_page = []
                text += para.text + "\n"
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    text += row_text + "\n"
                    current_page.append(row_text)
            
            if current_page:
                pages.append({"page": len(pages) + 1, "text": "\n".join(current_page)})
        except Exception as e:
            text = f"Word解析错误: {str(e)}"
        return text, pages

    def _process_xlsx(self, file_path: str) -> tuple:
        """解析Excel文件"""
        text = ""
        pages = []
        try:
            wb = load_workbook(file_path, read_only=True, data_only=True)
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                sheet_text = f"=== 工作表: {sheet_name} ===\n"
                for row in ws.iter_rows(values_only=True):
                    row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                    if row_text.strip(" |"):
                        sheet_text += row_text + "\n"
                text += sheet_text + "\n"
                pages.append({"page": len(pages) + 1, "text": sheet_text, "sheet": sheet_name})
            wb.close()
        except Exception as e:
            text = f"Excel解析错误: {str(e)}"
        return text, pages

    def _process_text(self, file_path: str) -> tuple:
        """解析纯文本文件"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="gbk") as f:
                text = f.read()
        pages = [{"page": 1, "text": text}]
        return text, pages

    def chunk_text(self, text: str, source: str = "", chunk_size: int = 500, overlap: int = 50) -> List[Dict]:
        """将文本分块，用于RAG检索"""
        if not text or not text.strip():
            return []

        # 清理文本
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = text.strip()

        chunks = []
        sentences = re.split(r'(?<=[。！？.!?;\n])', text)
        
        current_chunk = ""
        chunk_idx = 0
        
        for sentence in sentences:
            if not sentence.strip():
                continue
            
            if len(current_chunk) + len(sentence) <= chunk_size:
                current_chunk += sentence
            else:
                if current_chunk.strip():
                    chunks.append({
                        "id": f"{source}_chunk_{chunk_idx}",
                        "text": current_chunk.strip(),
                        "source": source,
                        "chunk_index": chunk_idx,
                        "char_count": len(current_chunk.strip())
                    })
                    chunk_idx += 1
                    # 保留overlap
                    current_chunk = current_chunk[-overlap:] + sentence if overlap > 0 else sentence
                else:
                    current_chunk = sentence
        
        # 最后一块
        if current_chunk.strip():
            chunks.append({
                "id": f"{source}_chunk_{chunk_idx}",
                "text": current_chunk.strip(),
                "source": source,
                "chunk_index": chunk_idx,
                "char_count": len(current_chunk.strip())
            })
        
        return chunks


# 全局文档处理器实例
doc_processor = DocumentProcessor()
